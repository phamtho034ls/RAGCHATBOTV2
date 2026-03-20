"""DOCX/DOC parser for Vietnamese legal documents.

Extracts raw text from .docx files (via python-docx) and .doc files
(via Microsoft Word COM automation on Windows).
"""

from __future__ import annotations

import logging
import platform
from pathlib import Path
from docx import Document as DocxDocument

log = logging.getLogger(__name__)


def read_docx(file_path: str | Path) -> str:
    """Extract full text from a DOCX or DOC file.

    - .docx: uses python-docx (cross-platform)
    - .doc:  uses Word COM automation (Windows only)
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if file_path.suffix.lower() == ".doc":
        return _parse_doc_com(file_path)

    return _parse_docx(file_path)


def parse_docx(file_path: str | Path) -> str:
    """Backward-compatible alias for legacy imports."""
    return read_docx(file_path)


def _parse_docx(file_path: Path) -> str:
    """Parse .docx via python-docx."""
    doc = DocxDocument(str(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    log.info("Parsed DOCX '%s': %d characters.", file_path.name, len(full_text))
    return full_text


def _parse_doc_com(file_path: Path) -> str:
    """Parse .doc via Microsoft Word COM automation (Windows only)."""
    if platform.system() != "Windows":
        raise RuntimeError(
            "Parsing .doc files requires Microsoft Word on Windows. "
            "Please convert to .docx first."
        )

    import win32com.client
    import pythoncom

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(file_path.resolve()))
        full_text = doc.Content.Text
        log.info("Parsed DOC '%s' via COM: %d characters.", file_path.name, len(full_text))
        return full_text
    finally:
        if doc:
            doc.Close(False)
        if word:
            word.Quit()
        pythoncom.CoUninitialize()
